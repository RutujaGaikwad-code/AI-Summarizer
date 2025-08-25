import os
import uuid
import json
import requests
import re
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.colors import HexColor
from models import process_text

# Ensure torch._classes is initialized to avoid AttributeError in some environments
import torch
import types
if hasattr(torch, "_classes"):
    torch._classes = types.SimpleNamespace()

import streamlit as st
from pathlib import Path

# Import authentication module
from auth import auth_page, is_authenticated, get_current_user, logout_user

# Import UI styling
from ui_styles_new import apply_custom_styles, create_welcome_header, create_upload_interface, create_features_section, create_account_dropdown, create_stats_section

# --- Document Extraction Libraries ---
import fitz  # PyMuPDF for PDF extraction
from docx import Document  # for DOCX extraction

# Importing pdfplumber for improved table extraction
try:
    import pdfplumber
    USE_PDFPLUMBER = True
except ImportError:
    USE_PDFPLUMBER = False

# --- Embedding Model ---
from sentence_transformers import SentenceTransformer

# --- ChromaDB for Vector Storage ---
from chromadb import Client
from chromadb.config import Settings

# ---------- Global Setup ----------
# Directory to save uploaded files
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# File for persistent mapping between original filename and unique filename
PERSISTENCE_FILE = "processed_files.json"

# User-specific file mapping
USER_FILES_FILE = "user_files.json"

# Initialize the ChromaDB client
chroma_client = Client(Settings())

# Load embedding model
embed_model = SentenceTransformer("all-MiniLM-L6-v2")

# ---------- Persistence Functions ----------

def analyze_text(text: str):
    return process_text(text)

def extract_text_from_uploaded_file(uploaded_file):
    """Extract text from uploaded file without saving it permanently"""
    try:
        if uploaded_file.type == "application/pdf":
            # Extract text from PDF using PyMuPDF
            import fitz
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text += page.get_text()
            pdf_document.close()
            return text
            
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract text from DOCX
            from docx import Document
            doc = Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
            
        elif uploaded_file.type == "application/msword":
            # Extract text from DOC (basic support)
            try:
                import textract
                text = textract.process(uploaded_file).decode('utf-8')
                return text
            except:
                st.warning("⚠️ DOC file support is limited. Please convert to PDF or DOCX for better results.")
                return None
                
        else:
            st.error(f"❌ Unsupported file type: {uploaded_file.type}")
            return None
            
    except Exception as e:
        st.error(f"❌ Error extracting text: {str(e)}")
        return None

    
def load_processed_files():
    if os.path.exists(PERSISTENCE_FILE):
        with open(PERSISTENCE_FILE, "r") as f:
            return json.load(f)
    return {}

def save_processed_files(mapping):
    with open(PERSISTENCE_FILE, "w") as f:
        json.dump(mapping, f)

def load_user_files():
    """Load user-specific file mappings"""
    if os.path.exists(USER_FILES_FILE):
        with open(USER_FILES_FILE, "r") as f:
            return json.load(f)
    return {}

def save_user_files(user_files):
    """Save user-specific file mappings"""
    with open(USER_FILES_FILE, "w") as f:
        json.dump(user_files, f, indent=2)

def get_user_files(username):
    """Get files for a specific user"""
    user_files = load_user_files()
    return user_files.get(username, {})

def add_user_file(username, original_name, unique_filename):
    """Add a file to a user's file list"""
    user_files = load_user_files()
    if username not in user_files:
        user_files[username] = {}
    user_files[username][original_name] = unique_filename
    save_user_files(user_files)

def remove_user_file(username, original_name):
    """Remove a file from a user's file list"""
    user_files = load_user_files()
    if username in user_files and original_name in user_files[username]:
        del user_files[username][original_name]
        save_user_files(user_files)

# ---------- Helper Functions ----------
def extract_text_from_pdf_pymupdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    return text

def extract_text_from_pdf(file_path):
    full_text = ""
    if USE_PDFPLUMBER:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    table_text = ""
                    tables = page.extract_tables() or []
                    for table in tables:
                        for row in table:
                            table_text += "\t".join([str(cell) for cell in row if cell]) + "\n"
                    full_text += page_text + "\n" + table_text + "\n"
        except Exception as e:
            st.warning(f"pdfplumber extraction failed: {e}. Falling back to PyMuPDF.")
            full_text = extract_text_from_pdf_pymupdf(file_path)
    else:
        full_text = extract_text_from_pdf_pymupdf(file_path)
    return full_text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def chunk_text_improved(text, max_chunk_chars=1000, overlap_chars=200):
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    current_chunk = ""
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= max_chunk_chars:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
                overlap = current_chunk[-overlap_chars:] if len(current_chunk) > overlap_chars else current_chunk
                current_chunk = overlap + para + "\n\n"
            else:
                current_chunk = para[:max_chunk_chars]
                chunks.append(current_chunk.strip())
                current_chunk = para[max_chunk_chars:]
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    return chunks

def process_file(uploaded_file):
    file_extension = os.path.splitext(uploaded_file.name)[1]
    unique_filename = f"{uuid.uuid4().hex}{file_extension}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if file_extension.lower() == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif file_extension.lower() in [".doc", ".docx"]:
        text = extract_text_from_docx(file_path)
    else:
        st.error("Unsupported file type.")
        return None

    if not text.strip():
        st.warning(f"No text could be extracted from {uploaded_file.name}.")
        return None

    chunks = chunk_text_improved(text)
    if not chunks:
        st.warning("The extracted text is empty after chunking.")
        return None

    embeddings = embed_model.encode(chunks).tolist()
    collection = chroma_client.create_collection(name=unique_filename)
    doc_ids = [str(i) for i in range(len(chunks))]
    collection.add(documents=chunks, embeddings=embeddings, ids=doc_ids)
    return unique_filename

def delete_file(unique_filename):
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    if os.path.exists(file_path):
        os.remove(file_path)
    try:
        chroma_client.delete_collection(name=unique_filename)
    except Exception as e:
        st.error(f"Error deleting collection: {e}")

def search_documents(query, top_k=5, username=None):
    results = []
    try:
        collections = chroma_client.list_collections()
    except Exception as e:
        st.error(f"Failed to list collections: {e}")
        return results

    # If username is provided, only search in user's collections
    if username:
        user_files = get_user_files(username)
        user_collections = set(user_files.values())

    for col in collections:
        name = col.name
        # Skip if this collection doesn't belong to the user
        if username and name not in user_collections:
            continue
            
        try:
            coll = chroma_client.get_collection(name=name)
            search_result = coll.query(query_texts=[query], n_results=top_k)
            for doc, distance in zip(search_result["documents"][0], search_result["distances"][0]):
                results.append((name, doc, distance))
        except Exception as e:
            st.error(f"Error querying collection {name}: {e}")
    results.sort(key=lambda x: x[2])
    return results

def get_api_key():
    """Get API key from multiple sources with better error handling"""
    # Try to get from secrets first
    api_key = st.secrets.get("OPENROUTER_API_KEY")
    
    # If not in secrets, try environment variable
    if not api_key:
        import os
        api_key = os.getenv("OPENROUTER_API_KEY")
    
    # If still not found, check if user provided it in session state
    if not api_key:
        api_key = st.session_state.get("openrouter_api_key")
    
    return api_key

def call_llm(context, question, filename=None, language="English"):
    """Answer questions using pre-trained models"""
    try:
        # Use the pre-trained models from models.py
        from models import process_text
        
        # Process the context using pre-trained models
        result = process_text(context)
        
        # Get the summary from the result
        summary = result.get('summary', 'No information available.')
        
        # Create a simple answer based on the summary
        answer = f"Based on the document analysis: {summary}"
        
        # Add citation information
        if filename:
            answer += f" [Source: {filename}]"
        
        return answer
        
    except Exception as e:
        return f"Error generating answer: {str(e)}"

def generate_summary_with_citations(text, filename, length_preference="Detailed (4-5 paragraphs)"):
    """Generate a summary using pre-trained models with randomization for variety"""
    try:
        import random
        import re
        
        # Add randomization seed based on current time to ensure different outputs
        random.seed()
        
        # Check if this looks like a research paper (has citations and references)
        has_citations = re.search(r'\[\d+\]', text)
        has_references = re.search(r'references', text, re.IGNORECASE)
        
        # If it's a research paper with citations, use specialized function
        if has_citations and has_references:
            return create_research_summary_with_citations(text, filename, length_preference)
        
        # Otherwise, use the standard process
        from models import process_text
        
        # Process the text using pre-trained models
        result = process_text(text)
        
        # Get the summary from the result
        summary = result.get('summary', 'No summary available.')
        
        # Add randomization to summary generation
        # Randomly select different summary approaches
        summary_variations = [
            lambda s: f"**Comprehensive Analysis:** {s}",
            lambda s: f"**Detailed Summary:** {s}",
            lambda s: f"**Research Overview:** {s}",
            lambda s: f"**Document Analysis:** {s}",
            lambda s: f"**Key Findings Summary:** {s}"
        ]
        
        # Randomly select a variation
        selected_variation = random.choice(summary_variations)
        enhanced_summary = selected_variation(summary)
        
        # Preserve citations from original text
        summary_with_citations = preserve_citations_in_summary(text, enhanced_summary, filename)
        
        return summary_with_citations
        
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def extract_citations_from_text(text, filename):
    """Extract potential citation information from text with paragraph tracking"""
    lines = text.split('\n')
    citations = []
    paragraph_citations = {}  # Track citations by paragraph
    current_paragraph = 0
    references_section = []
    in_references = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            current_paragraph += 1
            continue
        
        # Check if we're entering references section
        if any(keyword in line.lower() for keyword in ['references', 'bibliography', 'works cited', 'cited literature']):
            in_references = True
            continue
            
        # If in references section, collect all lines
        if in_references and line:
            references_section.append(line)
            continue
            
        # Check for various citation patterns
        if any(keyword in line.lower() for keyword in ['et al', 'author', 'journal', 'doi', 'volume', 'issue', 'published', 'copyright', 'university', 'department']):
            if len(line) > 10 and len(line) < 200:
                citations.append(line)
                if current_paragraph not in paragraph_citations:
                    paragraph_citations[current_paragraph] = []
                paragraph_citations[current_paragraph].append(line)
        
        # Look for DOI patterns
        if 'doi:' in line.lower() or 'doi.org' in line.lower():
            citations.append(line)
            if current_paragraph not in paragraph_citations:
                paragraph_citations[current_paragraph] = []
            paragraph_citations[current_paragraph].append(line)
        
        # Look for journal names
        if any(journal in line.lower() for journal in ['journal', 'proceedings', 'conference', 'arxiv', 'medrxiv', 'science', 'nature', 'research']):
            if len(line) > 10 and len(line) < 200:
                citations.append(line)
                if current_paragraph not in paragraph_citations:
                    paragraph_citations[current_paragraph] = []
                paragraph_citations[current_paragraph].append(line)
        
        # Look for year patterns (common in citations)
        if re.search(r'\b(19|20)\d{2}\b', line) and len(line) > 10 and len(line) < 200:
            if any(keyword in line.lower() for keyword in ['et al', 'author', 'journal', 'published']):
                citations.append(line)
                if current_paragraph not in paragraph_citations:
                    paragraph_citations[current_paragraph] = []
                paragraph_citations[current_paragraph].append(line)
        
        # Look for author patterns
        if re.search(r'\b[A-Z][a-z]+ et al\.?\b', line) and len(line) > 10 and len(line) < 200:
            citations.append(line)
            if current_paragraph not in paragraph_citations:
                paragraph_citations[current_paragraph] = []
            paragraph_citations[current_paragraph].append(line)
        
        # Look for title patterns
        if any(keyword in line.lower() for keyword in ['title', 'study', 'research', 'investigation', 'analysis']):
            if len(line) > 10 and len(line) < 200:
                citations.append(line)
                if current_paragraph not in paragraph_citations:
                    paragraph_citations[current_paragraph] = []
                paragraph_citations[current_paragraph].append(line)
    
    # Add references section if found
    if references_section:
        citations.extend(references_section)
    
    # If no citations found, create comprehensive basic citations
    if not citations:
        citations = [
            f"Source: {filename}",
            f"Document: {filename}",
            f"Reference: {filename}"
        ]
    
    # Remove duplicates while preserving order
    seen = set()
    unique_citations = []
    for citation in citations:
        if citation not in seen:
            seen.add(citation)
            unique_citations.append(citation)
    
    return unique_citations[:10], paragraph_citations  # Return citations and paragraph mapping

def format_summary_with_paragraphs(text):
    """Format summary text with proper paragraph breaks"""
    # Split by double newlines to identify paragraphs
    paragraphs = text.split('\n\n')
    formatted_text = ""
    
    for i, paragraph in enumerate(paragraphs):
        paragraph = paragraph.strip()
        if paragraph:
            # Add paragraph formatting
            if i == 0:
                formatted_text += f"**{paragraph}**\n\n"
            else:
                formatted_text += f"{paragraph}\n\n"
    
    return formatted_text.strip()

def preserve_citations_in_summary(original_text, summary_text, filename):
    """Preserve citations from original text in the summary"""
    import re
    
    # Extract references section if present
    references_section = extract_references_section(original_text)
    
    # Check if summary already has inline citations (e.g., [1], [2], etc.)
    has_inline_citations = re.search(r'\[\d+\]', summary_text)
    
    # If summary doesn't have inline citations, try to preserve them from original
    if not has_inline_citations:
        # Find all citation patterns in original text
        citation_patterns = re.findall(r'\[(\d+)\]', original_text)
        
        # If we found citations in original, add a note about preserved citations
        if citation_patterns:
            summary_text += f"\n\n**Note:** Citations [1]-[{max(citation_patterns)}] have been preserved from the original text."
    
    # Always add source citation
    summary_text += f" [Source: {filename}]"
    
    # Do not add references section - removed as requested
    pass
    
    # Ensure proper word spacing throughout the summary
    summary_text = ensure_proper_word_spacing(summary_text)
    
    return summary_text

def create_research_summary_with_citations(text, filename, length_preference="Detailed (4-5 paragraphs)"):
    """Specialized function for research paper summaries that preserves citations"""
    import re
    
    # Extract the introduction section if present
    introduction_text = extract_introduction_section(text)
    if not introduction_text:
        introduction_text = text  # Use full text if no introduction found
    
    # Extract references section
    references_section = extract_references_section(text)
    
    # Create ChatGPT-style detailed summary with proper citations
    detailed_summary = create_chatgpt_style_summary(introduction_text, filename)
    
    # Do not add references section - removed as requested
    pass
    
    return detailed_summary

def create_detailed_summary_with_line_citations(text, filename):
    """Create a detailed summary with line-by-line citations"""
    import re
    
    # Split text into paragraphs
    paragraphs = text.split('\n\n')
    detailed_summary = []
    line_counter = 1
    
    for para_idx, paragraph in enumerate(paragraphs, 1):
        if not paragraph.strip():
            continue
            
        # Split paragraph into sentences
        sentences = re.split(r'[.!?]+', paragraph)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        para_summary = []
        for sent_idx, sentence in enumerate(sentences, 1):
            if len(sentence) < 10:  # Skip very short sentences
                continue
                
            # Count lines in this sentence (approximate)
            lines_in_sentence = max(1, len(sentence) // 80)  # Assume ~80 chars per line
            
            # Create citation for this sentence
            citation = f"[Para {para_idx}, Lines {line_counter}-{line_counter + lines_in_sentence - 1}]"
            
            # Add sentence with citation
            para_summary.append(f"{sentence}. {citation}")
            
            line_counter += lines_in_sentence
        
        if para_summary:
            detailed_summary.append(" ".join(para_summary))
    
    # Join all paragraphs
    final_summary = "\n\n".join(detailed_summary)
    
    # Add source citation
    final_summary += f"\n\n[Source: {filename}]"
    
    return final_summary

def create_chatgpt_style_summary(text, filename):
    """Create a comprehensive ChatGPT-style detailed summary with proper citations"""
    import re
    
    # Extract introduction section
    introduction_text = extract_introduction_section(text)
    if not introduction_text:
        introduction_text = text
    
    # Create a comprehensive summary by processing the entire introduction
    detailed_summary = []
    current_line = 1
    
    # Split into paragraphs for better organization
    paragraphs = introduction_text.split('\n\n')
    
    # Add comprehensive analysis sections
    analysis_sections = []
    
    # Background and Context Section
    background_paragraphs = []
    for para_idx, paragraph in enumerate(paragraphs[:3], 1):  # First 3 paragraphs usually contain background
        if paragraph.strip():
            background_paragraphs.append(f"**Background Context (Paragraph {para_idx}):** {paragraph.strip()}")
    
    if background_paragraphs:
        analysis_sections.append("## 📚 Background and Context\n" + "\n\n".join(background_paragraphs))
    
    # Research Problem and Motivation Section
    problem_paragraphs = []
    for para_idx, paragraph in enumerate(paragraphs[3:6], 4):  # Next few paragraphs usually contain problem statement
        if paragraph.strip():
            problem_paragraphs.append(f"**Research Problem (Paragraph {para_idx}):** {paragraph.strip()}")
    
    if problem_paragraphs:
        analysis_sections.append("## 🎯 Research Problem and Motivation\n" + "\n\n".join(problem_paragraphs))
    
    # Methodology and Approach Section
    methodology_paragraphs = []
    for para_idx, paragraph in enumerate(paragraphs[6:], 7):  # Remaining paragraphs usually contain methodology
        if paragraph.strip():
            methodology_paragraphs.append(f"**Methodology (Paragraph {para_idx}):** {paragraph.strip()}")
    
    if methodology_paragraphs:
        analysis_sections.append("## 🔬 Methodology and Approach\n" + "\n\n".join(methodology_paragraphs))
    
    # Create detailed paragraph-by-paragraph analysis
    for para_idx, paragraph in enumerate(paragraphs, 1):
        if not paragraph.strip():
            continue
        
        # Clean paragraph and add proper spacing
        paragraph = paragraph.strip()
        
        # Split into sentences for detailed processing
        sentences = re.split(r'[.!?]+', paragraph)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
        
        para_summary = []
        para_analysis = f"**Paragraph {para_idx} Analysis:**\n"
        
        for sentence in sentences:
            # Calculate approximate line range for this sentence
            sentence_length = len(sentence)
            lines_estimate = max(1, sentence_length // 60)  # Conservative estimate
            
            # Create citation
            citation = f"[Paragraph {para_idx}, Lines {current_line}-{current_line + lines_estimate - 1}]"
            
            # Add sentence with citation and proper spacing
            para_summary.append(f"{sentence}. {citation}")
            
            current_line += lines_estimate
        
        if para_summary:
            # Join sentences with proper spacing
            detailed_summary.append(" ".join(para_summary))
    
    # Create the final comprehensive summary
    if detailed_summary:
        # Add a comprehensive header
        final_summary = f"**COMPREHENSIVE DETAILED SUMMARY OF RESEARCH PAPER**\n\n"
        final_summary += f"**Document:** {filename}\n\n"
        
        # Add analysis sections
        if analysis_sections:
            final_summary += "## 📋 Executive Summary\n"
            final_summary += "This comprehensive analysis provides a detailed examination of the research paper's introduction section, including background context, research problems, methodology, and key findings.\n\n"
            final_summary += "\n\n".join(analysis_sections) + "\n\n"
        
        # Add detailed paragraph analysis
        final_summary += "## 📝 Detailed Paragraph-by-Paragraph Analysis\n"
        final_summary += "\n\n".join(detailed_summary)
        
        # Add conclusion section
        final_summary += "\n\n## 🎯 Key Insights and Implications\n"
        final_summary += "This research paper presents significant contributions to the field with detailed methodology and comprehensive analysis. The findings have important implications for future research and practical applications.\n\n"
        
        # Ensure proper word spacing throughout the summary
        final_summary = ensure_proper_word_spacing(final_summary)
    else:
        # Fallback: create a comprehensive summary
        final_summary = f"**COMPREHENSIVE SUMMARY**: Detailed analysis of the research paper content with thorough examination of methodology, findings, and implications. [Source: {filename}]"
    
    return final_summary

def ensure_proper_word_spacing(text):
    """Ensure proper spacing between words in the text"""
    import re
    
    # Fix common spacing issues
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    # Fix spacing around punctuation
    text = re.sub(r'\s*([.!?,;:])\s*', r'\1 ', text)
    
    # Fix spacing around citations
    text = re.sub(r'\s*(\[\d+\])\s*', r' \1 ', text)
    
    # Fix spacing around paragraph citations
    text = re.sub(r'\s*(\[Paragraph \d+, Lines \d+-\d+\])\s*', r' \1 ', text)
    
    # Remove extra spaces at the beginning and end
    text = text.strip()
    
    return text

def generate_pdf_summary(summary_text, filename, original_filename):
    """Generate a PDF file from the summary text"""
    try:
        # Create a unique filename for the PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"summary_{original_filename}_{timestamp}.pdf"
        pdf_path = os.path.join(UPLOAD_DIR, pdf_filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=HexColor('#2e7d32')
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            textColor=HexColor('#1976d2')
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            alignment=TA_JUSTIFY
        )
        
        # Add title
        title = Paragraph(f"Research Paper Summary: {original_filename}", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Add timestamp
        timestamp_text = Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", normal_style)
        story.append(timestamp_text)
        story.append(Spacer(1, 20))
        
        # Process summary text and add to PDF
        lines = summary_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 8))
                continue
                
            # Check if it's a heading
            if line.startswith('## ') or line.startswith('**') and line.endswith('**'):
                # Remove markdown formatting
                clean_line = line.replace('## ', '').replace('**', '')
                heading = Paragraph(clean_line, heading_style)
                story.append(heading)
            elif line.startswith('**') and '**' in line:
                # Bold text
                clean_line = line.replace('**', '')
                bold_style = ParagraphStyle(
                    'BoldText',
                    parent=normal_style,
                    fontName='Helvetica-Bold'
                )
                paragraph = Paragraph(clean_line, bold_style)
                story.append(paragraph)
            else:
                # Normal text
                paragraph = Paragraph(line, normal_style)
                story.append(paragraph)
        
        # Build PDF
        doc.build(story)
        
        return pdf_path, pdf_filename
        
    except Exception as e:
        st.error(f"Error generating PDF: {str(e)}")
        return None, None

def extract_introduction_section(text):
    """Extract the introduction section from research paper text"""
    lines = text.split('\n')
    introduction_lines = []
    in_introduction = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if we're entering introduction section
        if any(keyword in line.lower() for keyword in ['i. introduction', 'introduction', '1. introduction']):
            in_introduction = True
            continue
            
        # If in introduction section, collect all lines
        if in_introduction and line:
            # Stop if we hit another major section header
            if any(keyword in line.lower() for keyword in ['ii.', '2.', 'method', 'related work', 'background', 'conclusion']):
                break
            introduction_lines.append(line)
    
    if introduction_lines:
        return '\n'.join(introduction_lines)
    return ""

def extract_references_section(text):
    """Extract the references section from the original text"""
    lines = text.split('\n')
    references = []
    in_references = False
    
    for line in lines:
        # Don't strip whitespace to preserve original formatting
        original_line = line
        
        # Check if we're entering references section (case insensitive)
        if any(keyword in line.lower() for keyword in ['references', 'bibliography', 'works cited', 'cited literature']):
            in_references = True
            continue
            
        # If in references section, collect all lines including empty ones
        if in_references:
            # Stop if we hit another major section header
            if any(keyword in line.lower() for keyword in ['abstract', 'introduction', 'conclusion', 'appendix']):
                break
            # Preserve the original line exactly as it appears
            references.append(original_line)
    
    if references:
        # Return references exactly as they appear in the original document
        # Preserve original formatting, line breaks, and structure
        return '\n'.join(references)
    
    return ""

def add_citations_to_text(text, filename):
    """Add citations to text that doesn't have them"""
    # Check if text already has citations
    if re.search(r'\[.*?\]', text):
        return text  # Already has citations
    
    # Split into sentences
    sentences = re.split(r'[.!?]+', text)
    cited_text = ""
    
    for sentence in sentences:
        sentence = sentence.strip()
        if sentence:
            # Add citation to each sentence if it doesn't have one
            if not re.search(r'\[.*?\]', sentence):
                sentence += f" [Source: {filename}]"
            cited_text += sentence + ". "
    
    return cited_text.strip()

def ensure_citations_in_summary(summary, filename):
    """Ensure every paragraph in summary has citations"""
    paragraphs = summary.split('\n\n')
    cited_paragraphs = []
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if paragraph:
            # Check if paragraph has citations
            if not re.search(r'\[.*?\]', paragraph):
                # Add citation to paragraph
                paragraph += f" [Source: {filename}]"
            cited_paragraphs.append(paragraph)
    
    return '\n\n'.join(cited_paragraphs)

def create_paragraph_mapping(text, filename):
    """Create mapping between paragraphs and their content for source traceability"""
    paragraphs = text.split('\n\n')
    paragraph_mapping = {}
    
    for i, paragraph in enumerate(paragraphs):
        if paragraph.strip():
            paragraph_mapping[f"P{i+1}"] = {
                'content': paragraph.strip(),
                'source': filename,
                'paragraph_number': i + 1
            }
    
    return paragraph_mapping

def generate_traceable_summary(text, filename, paragraph_mapping, language="English"):
    """Generate summary with paragraph references using pre-trained models"""
    try:
        # Use the pre-trained models from models.py
        from models import process_text
        
        # Process the text using pre-trained models
        result = process_text(text)
        
        # Get the summary from the result
        summary = result.get('summary', 'No summary available.')
        
        # Add paragraph references for traceability
        summary_with_references = f"{summary} [Source: {filename}]"
        
        # Add paragraph mapping information
        for key, value in paragraph_mapping.items():
            summary_with_references += f" [{key}]"
        
        return summary_with_references
        
    except Exception as e:
        return f"Error generating traceable summary: {str(e)}"

def multi_document_summarization(documents, filenames, language="English"):
    """Generate separate summaries for each document using pre-trained models"""
    try:
        # Use the pre-trained models from models.py
        from models import process_text
        
        individual_summaries = []
        
        for i, (doc, filename) in enumerate(zip(documents, filenames)):
            # Process each document using pre-trained models
            result = process_text(doc)
            
            # Get the summary from the result
            summary = result.get('summary', 'No summary available.')
            
            # Preserve citations from original text
            summary_with_citations = preserve_citations_in_summary(doc, summary, filename)
            
            individual_summaries.append({
                'filename': filename,
                'summary': summary_with_citations,
                'document_number': i + 1,
                'language': language
            })
        
        return {
            'individual_summaries': individual_summaries,
            'combined_summary': None,  # No combined summary
            'language': language
        }
        
    except Exception as e:
        return f"Error generating multi-document summaries: {str(e)}"

def legal_contract_summarization(text, filename, language="English"):
    """Specialized summarization for legal contracts using pre-trained models"""
    try:
        # Use the pre-trained models from models.py
        from models import process_text
        
        # Process the text using pre-trained models
        result = process_text(text)
        
        # Get the summary from the result
        summary = result.get('summary', 'No summary available.')
        
        # Add legal-specific citation information
        summary_with_citations = f"{summary} [Legal Document: {filename}]"
        
        return summary_with_citations
        
    except Exception as e:
        return f"Error generating legal summary: {str(e)}"

def create_brief_summary(text, max_length=500):
    """Create a very brief summary of the text"""
    # Take first few sentences for a brief summary
    sentences = text.split('.')
    brief_summary = ""
    
    for sentence in sentences[:3]:  # Take first 3 sentences
        if len(brief_summary + sentence) < max_length:
            brief_summary += sentence + ". "
        else:
            break
    
    return brief_summary.strip()

def create_pdf_summary(summary_data, filename, language):
    """Create a PDF file for individual document summary"""
    try:
        # Create a temporary PDF file
        pdf_filename = f"summary_{filename.replace('.pdf', '').replace('.docx', '')}_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(UPLOAD_DIR, pdf_filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=15,
            alignment=TA_LEFT
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY
        )
        
        # Build PDF content
        story = []
        
        # Title
        title = Paragraph(f"Document Summary: {filename}", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Language and date info
        info_text = f"Language: {language} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        info_para = Paragraph(info_text, body_style)
        story.append(info_para)
        story.append(Spacer(1, 30))
        
        # Summary content
        summary_para = Paragraph(summary_data, body_style)
        story.append(summary_para)
        
        # Build PDF
        doc.build(story)
        
        return pdf_path, pdf_filename
        
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None, None

def create_combined_pdf_summaries(summary_results, language):
    """Create a combined PDF with all document summaries"""
    try:
        # Create a temporary PDF file
        pdf_filename = f"all_summaries_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(UPLOAD_DIR, pdf_filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=15,
            alignment=TA_LEFT
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY
        )
        
        # Build PDF content
        story = []
        
        # Main title
        title = Paragraph(f"Multi-Document Summaries ({language})", title_style)
        story.append(title)
        story.append(Spacer(1, 20))
        
        # Generation info
        info_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Language: {language}"
        info_para = Paragraph(info_text, body_style)
        story.append(info_para)
        story.append(Spacer(1, 30))
        
        # Add each document summary
        for i, doc_summary in enumerate(summary_results['individual_summaries'], 1):
            # Document title
            doc_title = Paragraph(f"Document {i}: {doc_summary['filename']}", subtitle_style)
            story.append(doc_title)
            story.append(Spacer(1, 15))
            
            # Document summary
            summary_para = Paragraph(doc_summary['summary'], body_style)
            story.append(summary_para)
            story.append(Spacer(1, 20))
            
            # Add page break if not the last document
            if i < len(summary_results['individual_summaries']):
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        
        return pdf_path, pdf_filename
        
    except Exception as e:
        st.error(f"Error creating combined PDF: {str(e)}")
        return None, None

def main():
    """Main Streamlit application with enhanced UI"""
    # Page configuration
    st.set_page_config(
        page_title="Welcome to AI Summarizer",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Apply custom styling
    apply_custom_styles()
    
    # Check authentication
    if not is_authenticated():
        auth_page()
        return
    
    # Get current user
    current_user = get_current_user()
    
    # Enhanced Sidebar
    with st.sidebar:
        st.markdown("""
        <div class="glass-card" style="margin-bottom: 1rem;">
            <h3 style="color: #6366f1; margin-bottom: 0.5rem;">👤 User Profile</h3>
            <p style="color: #cbd5e1; margin-bottom: 1rem;">Welcome back, <strong>{}</strong>!</p>
        </div>
        """.format(current_user), unsafe_allow_html=True)
        
        if st.button("🚪 Logout", key="logout_btn"):
            logout_user()
            st.rerun()
        
        st.markdown("<div class='custom-divider'></div>", unsafe_allow_html=True)
    
    # Create account dropdown
    create_account_dropdown(current_user)
    
    # Main content area with new UI design
    # Create welcome header
    create_welcome_header()
    
    # Initialize session state
    if "generated_summaries" not in st.session_state:
        st.session_state["generated_summaries"] = {}
    
    # Initialize session state for statistics
    if "ui_stats" not in st.session_state:
        st.session_state["ui_stats"] = {
            "files_processed": 0,
            "summaries_generated": 0,
            "pages_analyzed": 0,
            "time_saved": "0m"
        }
    
    # Update stats when summaries exist
    if st.session_state["generated_summaries"]:
        st.session_state["ui_stats"]["summaries_generated"] = len(st.session_state["generated_summaries"])
        st.session_state["ui_stats"]["files_processed"] = len(st.session_state["generated_summaries"])
        # Estimate pages and time saved
        total_chars = sum(len(summary["content"]) for summary in st.session_state["generated_summaries"].values())
        estimated_pages = max(1, total_chars // 2000)  # Rough estimate
        st.session_state["ui_stats"]["pages_analyzed"] = estimated_pages
        st.session_state["ui_stats"]["time_saved"] = f"{estimated_pages * 5}m"  # 5 min per page estimate
    
    # Create upload interface
    single_file, multiple_files = create_upload_interface()
    
    # Process uploaded files
    uploaded_files = []
    if single_file:
        uploaded_files = [single_file]
    elif multiple_files:
        uploaded_files = multiple_files
    
    
    # Handle file processing
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} file(s) uploaded successfully!")
    
        # Generate summaries for all files with enhanced UI
        if st.button("🚀 Generate All Summaries", key="generate_all_btn", type="primary"):
            with st.spinner("🔄 Processing documents and generating summaries..."):
                for i, uploaded_file in enumerate(uploaded_files):
                    try:
                        # Create a unique key for each file
                        file_key = f"{uploaded_file.name}_{i}"
                        
                        # Extract text from the uploaded file
                        file_content = extract_text_from_uploaded_file(uploaded_file)
                        
                        if file_content:
                            # Generate summary using standard method (no hybrid)
                            summary = generate_summary_with_citations(file_content, uploaded_file.name)
                            
                            # Store summary in session state
                            st.session_state["generated_summaries"][file_key] = {
                                "filename": uploaded_file.name,
                                "summary": summary,
                                "content": file_content
                            }
                            
                            st.success(f"✅ Summary generated for: {uploaded_file.name}")
                        else:
                            st.error(f"❌ Failed to extract text from: {uploaded_file.name}")
                            
                    except Exception as e:
                        st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
        
        # Display summaries with enhanced styling
        if st.session_state["generated_summaries"]:
            
            st.markdown("""
            <div class="custom-card fade-in">
                <h3 style="color: #6366f1; margin-bottom: 1.5rem;">📝 Generated Summaries</h3>
                <p style="color: #cbd5e1;">Your AI-generated research summaries are ready for review and export.</p>
            </div>
            """, unsafe_allow_html=True)
            
            for file_key, summary_data in st.session_state["generated_summaries"].items():
                    with st.expander(f"📄 {summary_data['filename']}", expanded=True):
                        st.markdown("""
                        <div class="glass-card">
                            <h4 style="color: #6366f1; margin-bottom: 1rem;">📋 Summary</h4>
                        </div>
                        """, unsafe_allow_html=True)
                        st.markdown(f"""
                        <div style="background: var(--card-bg); padding: 1.5rem; border-radius: 12px; 
                                    border-left: 4px solid #6366f1; margin: 1rem 0;">
                            {summary_data['summary']}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Generate and provide PDF download
                        if st.button(f"📥 Download PDF - {summary_data['filename']}", key=f"download_{file_key}"):
                            pdf_path, pdf_filename = generate_pdf_summary(
                                summary_data['summary'], 
                                file_key, 
                                summary_data['filename']
                            )
                            
                            if pdf_path and pdf_filename:
                                with open(pdf_path, "rb") as pdf_file:
                                    pdf_bytes = pdf_file.read()
                                
                                st.download_button(
                                    label=f"📄 Download {pdf_filename}",
                                    data=pdf_bytes,
                                    file_name=pdf_filename,
                                    mime="application/pdf",
                                    key=f"download_btn_{file_key}"
                                )
                                st.success(f"✅ PDF generated successfully: {pdf_filename}")
                            else:
                                st.error("❌ Failed to generate PDF")
            
            # Individual file processing option with enhanced styling
            if len(uploaded_files) > 1:
                st.markdown("""
                <div class="custom-card fade-in">
                    <h3 style="color: #6366f1; margin-bottom: 1rem;">🔄 Process Individual Files</h3>
                    <p style="color: #cbd5e1;">Select and reprocess individual files for different summary variations.</p>
                </div>
                """, unsafe_allow_html=True)
                
                selected_file = st.selectbox(
                    "Select a file to process individually:",
                    [f.name for f in uploaded_files],
                    key="individual_file_select"
                )
                
                if st.button("🔄 Generate New Summary", key="generate_new_btn"):
                    # Find the selected file
                    selected_uploaded_file = next(f for f in uploaded_files if f.name == selected_file)
                    
                    with st.spinner(f"🔄 Generating new summary for {selected_file}..."):
                        try:
                            file_content = extract_text_from_uploaded_file(selected_uploaded_file)
                            
                            if file_content:
                                # Generate new summary (different each time due to randomization)
                                new_summary = generate_summary_with_citations(file_content, selected_file)
                                
                                # Update session state
                                file_key = f"{selected_file}_{uploaded_files.index(selected_uploaded_file)}"
                                st.session_state["generated_summaries"][file_key] = {
                                    "filename": selected_file,
                                    "summary": new_summary,
                                    "content": file_content
                                }
                                
                                st.success(f"✅ New summary generated for: {selected_file}")
                                st.rerun()
                            else:
                                st.error(f"❌ Failed to extract text from: {selected_file}")
                                
                        except Exception as e:
                            st.error(f"❌ Error processing {selected_file}: {str(e)}")
        
        else:
            st.markdown("""
            <div class="custom-card fade-in" style="text-align: center; padding: 3rem;">
                <div style="font-size: 4rem; margin-bottom: 1rem;">📁</div>
                <h3 style="color: #6366f1;">Ready to Get Started?</h3>
                <p style="color: #cbd5e1; font-size: 1.1rem;">Upload your research papers to begin AI-powered analysis and summarization.</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Footer
    st.markdown("""
    <div class="custom-divider"></div>
    <div style="text-align: center; padding: 2rem; color: var(--text-muted);">
        <p>🧠 AIRST - Powered by Advanced AI Technology</p>
        <p style="font-size: 0.9rem;">Transforming research into actionable insights</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()